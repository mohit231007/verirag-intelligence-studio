"""Page-aware parsing and deterministic paragraph-preserving chunking."""

from __future__ import annotations

import csv
import io
import re
import unicodedata
import zipfile
from collections.abc import Iterable
from pathlib import PurePath

from docx import Document
from pypdf import PdfReader

from .models import DocumentChunk, IngestionResult, SourcePage
from .security import neutralize_spreadsheet_formula, sha256_bytes, validate_upload

_WHITESPACE = re.compile(r"[\t\x0b\x0c\r ]+")
_BLANK_LINES = re.compile(r"\n{3,}")
_SENTENCE_BOUNDARY = re.compile(r"(?<=[.!?])\s+")


class DocumentParseError(ValueError):
    """Raised for malformed, encrypted, or unreadable documents."""


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text.replace("\x00", ""))
    text = "\n".join(_WHITESPACE.sub(" ", line).strip() for line in text.splitlines())
    return _BLANK_LINES.sub("\n\n", text).strip()


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("Could not decode the text file")


def _parse_pdf(content: bytes, filename: str, max_pages: int) -> tuple[list[SourcePage], list[str]]:
    try:
        reader = PdfReader(io.BytesIO(content), strict=False)
    except Exception as exc:  # pypdf exposes several parser-specific exception classes
        raise DocumentParseError(f"{filename} is not a valid PDF") from exc
    if reader.is_encrypted:
        raise DocumentParseError(f"{filename} is encrypted and cannot be processed")
    if len(reader.pages) > max_pages:
        raise DocumentParseError(f"{filename} exceeds the {max_pages}-page limit")

    pages: list[SourcePage] = []
    warnings: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if text:
            pages.append(SourcePage(filename, number, text))
        else:
            warnings.append(f"Page {number} contained no extractable text")
    return pages, warnings


def _parse_docx(content: bytes, filename: str) -> list[SourcePage]:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            uncompressed = sum(member.file_size for member in archive.infolist())
            if uncompressed > 50 * 1024 * 1024:
                raise DocumentParseError(f"{filename} expands beyond the 50 MB safety limit")
        document = Document(io.BytesIO(content))
    except (zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise DocumentParseError(f"{filename} is not a valid DOCX file") from exc

    blocks: list[str] = []
    blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = normalize_text("\n\n".join(blocks))
    return [SourcePage(filename, 1, text)] if text else []


def _parse_csv(content: bytes, filename: str) -> list[SourcePage]:
    decoded = _decode_text(content)
    try:
        dialect = csv.Sniffer().sniff(decoded[:8_192], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows: list[str] = []
    for row_number, row in enumerate(csv.reader(io.StringIO(decoded), dialect), start=1):
        if row_number > 10_000:
            raise DocumentParseError(f"{filename} exceeds the 10,000-row CSV limit")
        safe_cells = [neutralize_spreadsheet_formula(cell.strip()) for cell in row]
        rows.append(f"Row {row_number}: " + " | ".join(safe_cells))
    text = normalize_text("\n".join(rows))
    return [SourcePage(filename, 1, text)] if text else []


def parse_document(content: bytes, filename: str, max_pages: int) -> tuple[list[SourcePage], list[str]]:
    extension = PurePath(filename).suffix.lower()
    if extension == ".pdf":
        return _parse_pdf(content, filename, max_pages)
    if extension == ".docx":
        return _parse_docx(content, filename), []
    if extension == ".csv":
        return _parse_csv(content, filename), []
    text = normalize_text(_decode_text(content))
    return ([SourcePage(filename, 1, text)] if text else []), []


def _units(text: str) -> list[str]:
    """Split on paragraphs first, then sentences for oversized paragraphs."""

    units: list[str] = []
    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        units.extend(part.strip() for part in _SENTENCE_BOUNDARY.split(paragraph) if part.strip())
    return units


def _tail_for_overlap(text: str, overlap_chars: int) -> str:
    if overlap_chars <= 0 or not text:
        return ""
    tail = text[-overlap_chars:]
    first_space = tail.find(" ")
    return tail[first_space + 1 :] if first_space >= 0 else tail


def chunk_page(page: SourcePage, document_hash: str, chunk_size: int, overlap: int) -> list[DocumentChunk]:
    """Create bounded chunks without dropping content or using random identifiers."""

    chunks: list[DocumentChunk] = []
    buffer = ""
    cursor = 0

    def emit(value: str) -> None:
        nonlocal cursor
        clean = value.strip()
        if not clean:
            return
        start = page.text.find(clean[: min(80, len(clean))], max(cursor - overlap, 0))
        start = max(start, 0)
        end = min(start + len(clean), len(page.text))
        index = len(chunks) + 1
        chunk_id = f"{document_hash[:12]}_p{page.page_number}_c{index}"
        chunks.append(
            DocumentChunk(
                chunk_id=chunk_id,
                document_hash=document_hash,
                source_doc=page.source_doc,
                page_number=page.page_number,
                text=clean,
                char_start=start,
                char_end=end,
            )
        )
        cursor = end

    bounded_units: list[str] = []
    for unit in _units(page.text):
        remaining = unit
        while len(remaining) > chunk_size:
            window = remaining[:chunk_size]
            split_at = window.rfind(" ")
            if split_at < chunk_size // 2:
                split_at = chunk_size
            bounded_units.append(remaining[:split_at].strip())
            remaining = remaining[split_at:].strip()
        if remaining:
            bounded_units.append(remaining)

    for unit in bounded_units:
        candidate = f"{buffer} {unit}".strip()
        if not buffer or len(candidate) <= chunk_size:
            buffer = candidate
            continue

        emit(buffer)
        overlap_text = _tail_for_overlap(buffer, overlap)
        allowed_overlap = max(0, chunk_size - len(unit) - 1)
        if len(overlap_text) > allowed_overlap:
            overlap_text = _tail_for_overlap(overlap_text, allowed_overlap)
        buffer = f"{overlap_text} {unit}".strip()
    emit(buffer)
    return chunks


def ingest_document(
    content: bytes,
    filename: str,
    *,
    max_file_bytes: int,
    max_pages: int,
    chunk_size: int,
    chunk_overlap: int,
) -> IngestionResult:
    safe_name = validate_upload(filename, content, max_file_bytes)
    document_hash = sha256_bytes(content)
    pages, warnings = parse_document(content, safe_name, max_pages)
    if not pages:
        raise DocumentParseError(f"No extractable text found in {safe_name}")
    chunks = tuple(
        chunk for page in pages for chunk in chunk_page(page, document_hash, chunk_size, chunk_overlap)
    )
    if not chunks:
        raise DocumentParseError(f"No usable text chunks could be produced from {safe_name}")
    return IngestionResult(document_hash, safe_name, len(pages), chunks, tuple(warnings))


def total_chunks(results: Iterable[IngestionResult]) -> int:
    return sum(len(result.chunks) for result in results)
